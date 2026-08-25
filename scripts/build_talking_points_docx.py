#!/usr/bin/env python3
"""Build the one-page call-center talking-points handout from Markdown."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "demo" / "CALL-CENTER-TALKING-POINTS.md"
OUTPUT = ROOT / "docs" / "demo" / "CALL-CENTER-TALKING-POINTS.docx"

INK = RGBColor(35, 35, 43)
DIM = RGBColor(90, 88, 99)
AMBER = RGBColor(185, 116, 28)
BLUE = RGBColor(47, 112, 155)
LINE = "D8D5CF"


def set_cell_border(paragraph, color: str = LINE) -> None:
    """Add a compact bottom rule to a paragraph."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_inline(paragraph, text: str, *, base_color: RGBColor = INK) -> None:
    """Render the small subset of Markdown used by the handout."""
    pieces = re.split(r"(\*\*.+?\*\*)", text)
    for piece in pieces:
        if not piece:
            continue
        bold = piece.startswith("**") and piece.endswith("**")
        value = piece[2:-2] if bold else piece
        run = paragraph.add_run(value)
        run.bold = bold
        run.font.name = "Arial"
        run.font.size = Pt(8.7)
        run.font.color.rgb = BLUE if bold else base_color


def markdown_blocks(text: str):
    """Yield (kind, text) blocks while joining wrapped Markdown lines."""
    lines = text.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if not line.strip():
            index += 1
            continue
        if line.startswith("# "):
            yield "title", line[2:].strip()
            index += 1
            continue
        if line.startswith("## "):
            yield "heading", line[3:].strip()
            index += 1
            continue
        if line.startswith("- "):
            parts = [line[2:].strip()]
            index += 1
            while index < len(lines) and lines[index].startswith("  "):
                parts.append(lines[index].strip())
                index += 1
            yield "bullet", " ".join(parts)
            continue
        parts = [line.strip()]
        index += 1
        while index < len(lines):
            upcoming = lines[index]
            if not upcoming.strip() or upcoming.startswith("#") or upcoming.startswith("- "):
                break
            parts.append(upcoming.strip())
            index += 1
        yield "paragraph", " ".join(parts)


def build() -> None:
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.42)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.18)
    section.footer_distance = Inches(0.18)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(8.7)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(1.2)
    normal.paragraph_format.line_spacing = 1.0

    document.core_properties.title = "Cybic Voice Intelligence — Call Center Demonstration"
    document.core_properties.subject = "One-page call-center demonstration talking points"
    document.core_properties.author = "Cybic"

    for kind, text in markdown_blocks(SOURCE.read_text(encoding="utf-8")):
        if kind == "title":
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.keep_with_next = True
            run = paragraph.add_run(text)
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(15.5)
            run.font.color.rgb = AMBER
            set_cell_border(paragraph)
        elif kind == "heading":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(2.5)
            paragraph.paragraph_format.space_after = Pt(0.8)
            paragraph.paragraph_format.keep_with_next = True
            run = paragraph.add_run(text.upper())
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(9.4)
            run.font.color.rgb = AMBER
        elif kind == "bullet":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.15)
            paragraph.paragraph_format.first_line_indent = Inches(-0.12)
            paragraph.paragraph_format.space_after = Pt(0.4)
            bullet = paragraph.add_run("•  ")
            bullet.bold = True
            bullet.font.name = "Arial"
            bullet.font.size = Pt(8.7)
            bullet.font.color.rgb = AMBER
            add_inline(paragraph, text)
        else:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(1.2)
            add_inline(paragraph, text, base_color=DIM if text.startswith("The current demo validates") else INK)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("CYBIC  •  CALL CENTER TALKING POINTS")
    footer_run.font.name = "Arial"
    footer_run.font.size = Pt(6.5)
    footer_run.font.color.rgb = DIM

    document.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    try:
        build()
    except Exception as exc:  # pragma: no cover - command-line boundary
        print(f"error: {exc}", file=sys.stderr)
        raise
